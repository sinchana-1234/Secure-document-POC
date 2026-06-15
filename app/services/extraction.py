"""
extraction.py — turn an uploaded PDF / DOCX / TXT into clean plain text + metadata.

This module is intentionally strict. Extraction is where messy real-world files enter
the system, so instead of "open it and hope", every path is validated and every known
failure mode raises a SPECIFIC exception the upload layer can translate into a clear
message. Scope here is PDF, DOCX, TXT only — images and scanned PDFs are added in the
OCR module (the code is structured so that slots in cleanly).
"""
from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field
from typing import List, Optional

import pdfplumber
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from charset_normalizer import from_bytes

logger = logging.getLogger("doc-poc.extraction")


# ---------------------------------------------------------------------------
# Typed errors — the caller catches these to give the user a precise message.
# ---------------------------------------------------------------------------
class ExtractionError(Exception):
    """Base class for every extraction failure."""


class UnsupportedFileType(ExtractionError):
    """Extension/type isn't one this module handles."""


class EmptyFileError(ExtractionError):
    """The file on disk is 0 bytes."""


class CorruptFileError(ExtractionError):
    """File is unreadable, truncated, or its bytes don't match its extension."""


class EncryptedPDFError(ExtractionError):
    """PDF is password-protected; we can't read it without the password."""


class NoExtractableTextError(ExtractionError):
    """File opened fine but contains no usable text (e.g. a scanned PDF)."""


# ---------------------------------------------------------------------------
# Config / lookup tables
# ---------------------------------------------------------------------------
_EXT_TO_TYPE = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".text": "txt",
    ".md": "txt",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
}

_MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "image": "image/png",  # representative; actual subtype varies (png/jpg/webp)
}

_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = b"PK\x03\x04"          # .docx is a zip archive underneath
_SCANNED_PDF_MIN_CHARS = 10

SUPPORTED_EXTENSIONS = tuple(_EXT_TO_TYPE.keys())


@dataclass
class ExtractionResult:
    text: str
    doc_type: str                       # pdf | docx | txt
    mime_type: str
    page_count: Optional[int] = None    # PDFs only
    char_count: int = 0
    word_count: int = 0
    encoding: Optional[str] = None      # TXT only
    ocr_used: bool = False              # always False here; OCR module sets True later
    warnings: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def detect_doc_type(filename: str) -> str:
    """Map a filename to a normalized doc_type, or 'unknown' if unsupported."""
    ext = os.path.splitext(filename or "")[1].lower()
    return _EXT_TO_TYPE.get(ext, "unknown")


def extract(path: str, original_filename: str) -> ExtractionResult:
    """
    Validate and extract text from a supported file. Raises a typed ExtractionError
    on any problem so the caller can map it to a precise response.
    """
    doc_type = detect_doc_type(original_filename)
    if doc_type == "unknown":
        raise UnsupportedFileType(
            f"Unsupported file type '{os.path.splitext(original_filename)[1]}'. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}."
        )

    _validate_exists_and_nonempty(path)
    # Images skip the PDF/ZIP magic-byte check (different signatures); the vision
    # call will fail clearly if the bytes aren't a real image.
    if doc_type != "image":
        _validate_signature(path, doc_type)

    ocr_used = False
    if doc_type == "pdf":
        text, page_count = _extract_pdf(path)
        encoding = None
    elif doc_type == "docx":
        text, page_count, encoding = _extract_docx(path), None, None
    elif doc_type == "image":
        text = _extract_image_text(path)   # GPT-4o vision OCR
        page_count, encoding, ocr_used = None, None, True
    else:  # txt
        text, encoding = _extract_txt(path)
        page_count = None

    text = _normalize(text)
    if not text:
        if doc_type == "image":
            raise NoExtractableTextError(
                "No readable text was found in this image. Try a clearer photo or a file with text."
            )
        raise NoExtractableTextError("No readable text could be extracted from the file.")

    result = ExtractionResult(
        text=text,
        doc_type=doc_type,
        mime_type=_MIME[doc_type],
        page_count=page_count,
        char_count=len(text),
        word_count=len(text.split()),
        encoding=encoding,
        ocr_used=ocr_used,
    )
    logger.info("Extracted %s: %d chars, %d words%s",
                doc_type, result.char_count, result.word_count,
                f", {page_count} pages" if page_count else "")
    return result


# ---------------------------------------------------------------------------
# Validation helpers
# ---------------------------------------------------------------------------
def _validate_exists_and_nonempty(path: str) -> None:
    if not os.path.isfile(path):
        raise CorruptFileError(f"File not found on disk: {path}")
    if os.path.getsize(path) == 0:
        raise EmptyFileError("The uploaded file is empty (0 bytes).")


def _read_head(path: str, n: int = 8) -> bytes:
    with open(path, "rb") as f:
        return f.read(n)


def _validate_signature(path: str, doc_type: str) -> None:
    """Confirm the file's real bytes match its extension (catches renamed/wrong files)."""
    head = _read_head(path)
    if doc_type == "pdf" and not head.startswith(_PDF_MAGIC):
        raise CorruptFileError("File has a .pdf extension but is not a PDF (bad header).")
    if doc_type == "docx" and not head.startswith(_ZIP_MAGIC):
        raise CorruptFileError("File has a .docx extension but is not a valid Office file (bad header).")
    if doc_type == "txt":
        # UTF-16/UTF-32 are valid text but contain null bytes, so a BOM proves "text".
        text_boms = (b"\xef\xbb\xbf", b"\xff\xfe", b"\xfe\xff",
                     b"\xff\xfe\x00\x00", b"\x00\x00\xfe\xff")
        if not any(head.startswith(b) for b in text_boms) and b"\x00" in head:
            raise CorruptFileError("File has a text extension but appears to be binary.")


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------
def _extract_pdf(path: str):
    """Return (text, page_count). Extracts the text layer AND any embedded images via vision."""
    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            parts = [(page.extract_text() or "") for page in pdf.pages]
    except Exception as e:  # noqa: BLE001
        blob = []
        cur = e
        while cur is not None:
            blob.append(type(cur).__name__.lower())
            blob.append(str(cur).lower())
            cur = cur.__cause__ or cur.__context__
        joined = " ".join(blob)
        if "password" in joined or "encrypt" in joined:
            raise EncryptedPDFError("PDF is password-protected and cannot be read.") from e
        raise CorruptFileError(f"Could not read PDF (corrupt or malformed): {e}") from e

    text = "\n\n".join(p for p in parts if p.strip())

    # Extract text from images embedded in the PDF (charts, screenshots, scanned pages).
    image_texts = _extract_pdf_image_texts(path)
    if image_texts:
        text = (text + "\n\n" + "\n\n".join(image_texts)).strip()

    if len(text.strip()) < _SCANNED_PDF_MIN_CHARS:
        raise NoExtractableTextError(
            "PDF has no extractable text (no text layer and no readable images)."
        )
    return text, page_count


def _extract_pdf_image_texts(path: str) -> List[str]:
    """Pull every embedded image from the PDF and OCR each via vision. Per-image errors are skipped."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF (fitz) not installed; skipping embedded-image extraction in PDF.")
        return []

    results: List[str] = []
    try:
        pdf = fitz.open(path)
    except Exception as e:  # noqa: BLE001
        logger.warning("Could not open PDF for image extraction: %s", e)
        return results

    try:
        for page_index in range(len(pdf)):
            for img in pdf[page_index].get_images(full=True):
                xref = img[0]
                try:
                    base = pdf.extract_image(xref)
                    img_bytes = base["image"]
                    ext = base.get("ext", "png")
                    txt = _extract_image_text_from_bytes(img_bytes, media=ext)
                    if txt:
                        results.append(txt)
                except ExtractionError:
                    raise
                except Exception as e:  # noqa: BLE001
                    logger.warning("Skipped one embedded PDF image (page %d): %s", page_index, e)
                    continue
    finally:
        pdf.close()
    return results

def _extract_docx(path: str) -> str:
    """Return text from paragraphs + tables + any embedded images (via vision)."""
    try:
        doc = DocxDocument(path)
    except PackageNotFoundError as e:
        raise CorruptFileError("Not a valid .docx file (corrupt or wrong format).") from e
    except Exception as e:  # noqa: BLE001
        raise CorruptFileError(f"Could not read DOCX: {e}") from e

    parts: List[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    # Extract text from images embedded in the DOCX (screenshots, charts, scans).
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                ext = (rel.target_part.content_type or "").split("/")[-1] or "png"
                txt = _extract_image_text_from_bytes(img_bytes, media=ext)
                if txt:
                    parts.append(txt)
            except ExtractionError:
                raise
            except Exception as e:  # noqa: BLE001
                logger.warning("Skipped one embedded DOCX image: %s", e)
                continue

    return "\n".join(parts)

def _extract_txt(path: str):
    """
    Return (text, encoding). We try UTF-8 strictly first (covers the vast majority and
    is unambiguous); only if that fails do we fall back to detection for legacy encodings.
    """
    with open(path, "rb") as f:
        raw = f.read()
    if not raw.strip():
        raise EmptyFileError("The text file contains only whitespace.")

    try:
        return raw.decode("utf-8-sig"), "utf-8"
    except UnicodeDecodeError:
        pass
    best = from_bytes(raw).best()
    if best is None:
        return raw.decode("utf-8", errors="replace"), "utf-8 (fallback)"
    return str(best), best.encoding


# ---------------------------------------------------------------------------
# Image OCR via GPT-4o vision — reads text/information out of an uploaded image.
# ---------------------------------------------------------------------------
def _extract_image_text_from_bytes(image_bytes: bytes, media: str = "png") -> str:
    """
    Core vision OCR: send raw image bytes to GPT-4o vision and return extracted text.
    Shared by standalone image uploads AND images embedded in PDFs / DOCX.
    """
    import base64
    from openai import OpenAI, OpenAIError, AuthenticationError, RateLimitError
    from app.config import settings

    if not settings.OPENAI_API_KEY:
        raise ExtractionError("OPENAI_API_KEY is not set; image text extraction needs it.")

    b64 = base64.b64encode(image_bytes).decode("utf-8")
    media = "jpeg" if media in ("jpg", "jpeg") else media

    client = OpenAI(api_key=settings.OPENAI_API_KEY, max_retries=3, timeout=60.0)
    try:
        resp = client.chat.completions.create(
            model=settings.VISION_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text":
                        "Extract all text and information visible in this image. "
                        "Return only the extracted text, preserving structure where helpful. "
                        "If there is no text, reply with an empty response."},
                    {"type": "image_url",
                     "image_url": {"url": f"data:image/{media};base64,{b64}"}},
                ],
            }],
            max_tokens=2000,
        )
    except AuthenticationError as e:
        raise ExtractionError("OpenAI rejected the API key during image extraction.") from e
    except RateLimitError as e:
        raise ExtractionError("OpenAI rate limit hit during image extraction — try again later.") from e
    except OpenAIError as e:
        raise ExtractionError(f"Image text extraction failed: {e}") from e

    return (resp.choices[0].message.content or "").strip()


def _extract_image_text(path: str) -> str:
    """Standalone image upload: read the file and run vision OCR on its bytes."""
    with open(path, "rb") as f:
        data = f.read()
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    return _extract_image_text_from_bytes(data, media=ext)

# ---------------------------------------------------------------------------
# Output normalization
# ---------------------------------------------------------------------------
def _normalize(text: str) -> str:
    """Standardize line endings, trim trailing spaces, collapse runaway blank lines."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()